import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with fitz.open(pdf_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text += page.get_text()
    return text

def make_notes_in_word(text, word_path):
    """Create a Word document with extracted text as notes."""
    doc = Document()
    doc.add_heading('PDF Notes', level=1)
    paragraphs = text.split('\n')

    for paragraph in paragraphs:
        if paragraph.strip():  # Ignore empty lines
            doc.add_paragraph(paragraph)

    doc.save(word_path)

def main(pdf_path, word_path):
    text = extract_text_from_pdf(pdf_path)
    make_notes_in_word(text, word_path)
    print(f"Notes saved to {word_path}")

def require():
    pdf_path = input("PDF path:")
    word_path = input("WORD path")
    main(pdf_path, word_path)